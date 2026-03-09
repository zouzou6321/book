#!/usr/bin/env python3
"""
Novel reading & AI creation server.
Serves the static website and provides API endpoints for AI-powered novel creation.
"""

import base64
import json
import os
import subprocess
import re
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

from flask import Flask, request, Response, send_from_directory, jsonify
import requests as http_requests

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from consistency import run_check as run_consistency_check
import tasks as task_manager

app = Flask(__name__)

# IP 封禁：某 IP 累计 10 次收到 404 后永久封禁，之后对该 IP 始终返回 404
_BAN_LOCK = threading.Lock()
_BANNED_IPS = set()
_IP_404_COUNT = {}
_404_BAN_THRESHOLD = 10


def _client_ip():
    """优先取 X-Forwarded-For 首段（真实客户端），否则 remote_addr。"""
    xff = request.headers.get("X-Forwarded-For")
    if xff:
        return xff.split(",")[0].strip()
    return request.remote_addr or ""


@app.before_request
def _block_banned_ips():
    ip = _client_ip()
    if not ip:
        return None
    with _BAN_LOCK:
        if ip in _BANNED_IPS:
            return Response("Not Found", status=404, mimetype="text/plain")


@app.after_request
def _track_404_for_ban(response):
    if response.status_code != 404:
        return response
    ip = _client_ip()
    if not ip:
        return response
    with _BAN_LOCK:
        if ip in _BANNED_IPS:
            return response
        count = _IP_404_COUNT.get(ip, 0) + 1
        _IP_404_COUNT[ip] = count
        if count >= _404_BAN_THRESHOLD:
            _BANNED_IPS.add(ip)
    return response


BOOKS_ROOT = Path(__file__).resolve().parent.parent
WEBSITE_DIR = Path(__file__).resolve().parent
SKIP_DIRS = {"website", "novel-website", ".cursor", ".git", "node_modules"}

# Built-in AI (NVIDIA build.nvidia.com)
BUILTIN_API_KEY = "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"
BUILTIN_BASE_URL = "https://integrate.api.nvidia.com/v1"
BUILTIN_MODEL = "z-ai/glm4.7"

# NVIDIA FLUX image generation
FLUX_IMAGE_URL = "https://ai.api.nvidia.com/v1/genai/black-forest-labs/flux.1-dev"


def resolve_ai_config(data: dict) -> tuple:
    """Resolve API key, base URL, and model from request → env → built-in."""
    api_key = data.get("api_key") or os.getenv("AI_API_KEY") or BUILTIN_API_KEY
    base_url = (data.get("base_url") or os.getenv("AI_BASE_URL") or BUILTIN_BASE_URL).rstrip("/")
    model = data.get("model") or os.getenv("AI_MODEL") or BUILTIN_MODEL
    return api_key, base_url, model


# Simple Unicode → ASCII transliteration table for common CJK
_PINYIN_MAP = {
    "的": "de", "一": "yi", "是": "shi", "了": "le", "不": "bu",
    "在": "zai", "有": "you", "人": "ren", "这": "zhe", "中": "zhong",
    "大": "da", "天": "tian", "下": "xia", "上": "shang", "小": "xiao",
    "之": "zhi", "凤": "feng", "龙": "long", "谋": "mou", "星": "xing",
    "铸": "zhu", "者": "zhe", "世": "shi", "界": "jie", "战": "zhan",
    "神": "shen", "王": "wang", "剑": "jian", "魔": "mo", "仙": "xian",
    "道": "dao", "无": "wu", "风": "feng", "云": "yun", "火": "huo",
    "雷": "lei", "月": "yue", "日": "ri", "山": "shan", "海": "hai",
    "玄": "xuan", "血": "xue", "帝": "di", "皇": "huang", "霸": "ba",
    "修": "xiu", "破": "po", "九": "jiu", "万": "wan", "千": "qian",
    "百": "bai", "十": "shi", "古": "gu", "今": "jin", "城": "cheng",
    "梦": "meng", "情": "qing", "心": "xin", "灵": "ling", "诀": "jue",
    "传": "chuan", "记": "ji", "录": "lu", "经": "jing", "武": "wu",
    "侠": "xia", "影": "ying", "光": "guang", "夜": "ye", "暗": "an",
    "黑": "hei", "白": "bai", "红": "hong", "青": "qing", "金": "jin",
    "银": "yin", "玉": "yu", "花": "hua", "雪": "xue", "冰": "bing",
    "霜": "shuang", "绝": "jue", "逆": "ni", "乱": "luan", "斗": "dou",
    "杀": "sha", "生": "sheng", "死": "si", "命": "ming", "运": "yun",
    "图": "tu", "鉴": "jian", "域": "yu", "国": "guo", "族": "zu",
}


def _title_to_dir_id(title: str) -> str:
    """Convert a Chinese title to a filesystem-safe directory name."""
    parts = []
    ascii_buf = []

    def flush_ascii():
        if ascii_buf:
            parts.append(''.join(ascii_buf))
            ascii_buf.clear()

    for ch in title:
        if ch.isascii() and (ch.isalnum() or ch == '_'):
            ascii_buf.append(ch.lower())
        else:
            flush_ascii()
            if ch in _PINYIN_MAP:
                parts.append(_PINYIN_MAP[ch])
            elif '\u4e00' <= ch <= '\u9fff':
                parts.append(f"{ord(ch):x}")
    flush_ascii()

    result = '_'.join(parts) if parts else "novel"
    return result + "_novel"


# ──────────────────────────── Static Files ────────────────────────────

@app.route("/")
def index():
    return send_from_directory(WEBSITE_DIR, "index.html")


@app.route("/manifest.webmanifest")
def serve_manifest():
    return send_from_directory(WEBSITE_DIR, "manifest.webmanifest", mimetype="application/manifest+json")


@app.route("/sw.js")
def serve_sw():
    return send_from_directory(WEBSITE_DIR, "sw.js", mimetype="application/javascript")


@app.route("/<path:filepath>")
def static_files(filepath):
    return send_from_directory(WEBSITE_DIR, filepath)


# ──────────────────────────── AI Streaming ────────────────────────────

@app.route("/api/ai/config", methods=["GET"])
def ai_config():
    """Tell the frontend what AI is configured."""
    has_custom = bool(os.getenv("AI_API_KEY"))
    return jsonify({
        "has_builtin": bool(BUILTIN_API_KEY),
        "has_custom": has_custom,
        "builtin_model": BUILTIN_MODEL,
    })


@app.route("/api/ai/stream", methods=["POST"])
def ai_stream():
    data = request.json
    api_key, base_url, model = resolve_ai_config(data)
    messages = data.get("messages", [])

    if not api_key:
        return jsonify({"error": "未配置 API Key。请在设置中填写。"}), 400

    def generate():
        url = f"{base_url}/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": model,
            "messages": messages,
            "stream": True,
            "max_tokens": data.get("max_tokens", 16384),
        }

        in_think = False
        think_buf = ""

        try:
            with http_requests.post(url, json=payload, headers=headers, stream=True, timeout=180) as resp:
                if resp.status_code != 200:
                    error_text = resp.text[:500]
                    yield f"data: {json.dumps({'error': f'API 错误 ({resp.status_code}): {error_text}'})}\n\n"
                    return

                for line in resp.iter_lines(decode_unicode=True):
                    if not line or not line.startswith("data: "):
                        continue
                    chunk = line[6:].strip()
                    if chunk == "[DONE]":
                        yield "data: [DONE]\n\n"
                        return
                    try:
                        obj = json.loads(chunk)
                        delta = obj["choices"][0].get("delta", {})
                        content = delta.get("content", "")
                        if not content:
                            continue

                        # Filter <think>...</think> blocks from reasoning models
                        think_buf += content
                        while think_buf:
                            if in_think:
                                end_idx = think_buf.find("</think>")
                                if end_idx == -1:
                                    think_buf = ""
                                    break
                                in_think = False
                                think_buf = think_buf[end_idx + 8:]
                            else:
                                start_idx = think_buf.find("<think>")
                                if start_idx == -1:
                                    if think_buf:
                                        yield f"data: {json.dumps({'content': think_buf})}\n\n"
                                    think_buf = ""
                                    break
                                if start_idx > 0:
                                    yield f"data: {json.dumps({'content': think_buf[:start_idx]})}\n\n"
                                in_think = True
                                think_buf = think_buf[start_idx + 7:]
                    except (json.JSONDecodeError, KeyError, IndexError):
                        continue
        except http_requests.exceptions.ConnectionError:
            yield f"data: {json.dumps({'error': '无法连接到 AI 服务。请检查 Base URL 设置。'})}\n\n"
        except http_requests.exceptions.Timeout:
            yield f"data: {json.dumps({'error': 'AI 服务响应超时。'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ──────────────────────────── Novel Management ────────────────────────────

@app.route("/api/novel/list", methods=["GET"])
def novel_list():
    novels = []
    for d in sorted(BOOKS_ROOT.iterdir()):
        if not d.is_dir() or d.name in SKIP_DIRS or d.name.startswith("."):
            continue
        meta_file = d / "meta.json"
        if not meta_file.exists():
            continue
        meta = json.loads(meta_file.read_text(encoding="utf-8"))
        chapter_count = 0
        total_chars = 0
        for vol_dir in sorted(d.iterdir()):
            if vol_dir.is_dir() and vol_dir.name.startswith("volume_"):
                for cf in vol_dir.glob("chapter_*.md"):
                    chapter_count += 1
                    total_chars += len(cf.read_text(encoding="utf-8"))
        cover_exists = (d / "cover.jpg").exists()
        novels.append({
            "id": d.name,
            "meta": meta,
            "chapter_count": chapter_count,
            "total_chars": total_chars,
            "cover_image": cover_exists,
        })
    return jsonify(novels)


@app.route("/api/novel/init", methods=["POST"])
def novel_init():
    data = request.json
    meta = data["meta"]
    volume_count = data.get("volume_count", 3)
    total_chapters = data.get("total_chapters", 50)

    # Auto-generate directory ID from title via pinyin transliteration
    title = meta.get("title", "novel")
    novel_id = _title_to_dir_id(title)

    # Deduplicate if directory already exists
    base_id = novel_id
    counter = 2
    while (BOOKS_ROOT / novel_id).exists():
        novel_id = f"{base_id}_{counter}"
        counter += 1

    novel_dir = BOOKS_ROOT / novel_id

    novel_dir.mkdir(parents=True)

    # Distribute chapters evenly across volumes
    base = total_chapters // volume_count
    remainder = total_chapters % volume_count
    chapters_per_vol = [base + (1 if i < remainder else 0) for i in range(volume_count)]

    volumes = {}
    num_map = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    ch_counter = 1

    for i in range(volume_count):
        vol_dir_name = f"volume_{i + 1}"
        vol_display = f"第{num_map[i] if i < len(num_map) else i + 1}卷"
        volumes[vol_dir_name] = vol_display
        vol_path = novel_dir / vol_dir_name
        vol_path.mkdir()

        for _ in range(chapters_per_vol[i]):
            ch_file = vol_path / f"chapter_{ch_counter:03d}.md"
            ch_file.write_text(f"# 第{ch_counter:03d}章\n", encoding="utf-8")
            ch_counter += 1

    meta_data = {
        "title": meta["title"],
        "subtitle": meta.get("subtitle", ""),
        "author": meta.get("author", "AI & Thomas"),
        "description": meta.get("description", ""),
        "tags": meta.get("tags", []),
        "volumes": volumes,
    }
    if meta.get("premise") is not None:
        meta_data["premise"] = meta["premise"]
    if meta.get("writing_style"):
        meta_data["writing_style"] = meta["writing_style"]
    if meta.get("channel"):
        meta_data["channel"] = meta["channel"]
    (novel_dir / "meta.json").write_text(
        json.dumps(meta_data, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Initialize story bible and consistency directory
    (novel_dir / "story_bible.md").write_text(
        f"# {meta['title']} · 故事圣经\n\n"
        "> 本文件由系统自动维护，记录全书事实基准。AI 创作时必须严格遵守。\n\n"
        "## 角色登记\n\n| 角色 | 首次出现 | 身份 | 关键特征 |\n|------|---------|------|----------|\n\n"
        "## 时间线\n\n（暂无）\n\n"
        "## 关键事件\n\n（暂无）\n\n"
        "## 角色状态变更\n\n（暂无）\n",
        encoding="utf-8",
    )
    (novel_dir / "consistency").mkdir(exist_ok=True)

    return jsonify({"ok": True, "id": novel_id, "volumes": volumes})


@app.route("/api/novel/structure", methods=["POST"])
def novel_structure():
    data = request.json
    novel_id = data["novel_id"]
    novel_dir = BOOKS_ROOT / novel_id

    if not novel_dir.exists():
        return jsonify({"error": "Novel not found"}), 404

    meta = {}
    if (novel_dir / "meta.json").exists():
        meta = json.loads((novel_dir / "meta.json").read_text(encoding="utf-8"))

    structure = {
        "id": novel_id,
        "meta": meta,
        "has_outline": (novel_dir / "global_outline.md").exists(),
        "has_characters": (novel_dir / "characters.md").exists(),
        "volumes": [],
    }

    vol_dirs = sorted(
        [d for d in novel_dir.iterdir() if d.is_dir() and d.name.startswith("volume_")],
        key=_vol_sort_key,
    )

    for vol_dir in vol_dirs:
        vol_name = meta.get("volumes", {}).get(vol_dir.name, vol_dir.name)
        chapters = []
        for cf in sorted(vol_dir.glob("chapter_*.md"), key=_ch_sort_key):
            content = cf.read_text(encoding="utf-8")
            title = ""
            for line in content.split("\n"):
                if line.strip().startswith("# "):
                    title = line.strip()[2:]
                    break
            chapters.append({
                "id": cf.stem,
                "filename": cf.name,
                "title": title,
                "char_count": len(content),
            })

        structure["volumes"].append({
            "dir": vol_dir.name,
            "name": vol_name,
            "has_outline": (vol_dir / "outline_detailed.md").exists(),
            "chapters": chapters,
        })

    return jsonify(structure)


def _vol_sort_key(d):
    m = re.match(r"volume_(\d+)", d.name)
    return int(m.group(1)) if m else 0


def _ch_sort_key(cf):
    m = re.match(r"chapter_(\d+)", cf.stem)
    return int(m.group(1)) if m else 0


def _build_auto_steps(novel_dir: Path, meta: dict) -> list:
    """Build the same step list as frontend _buildAutoSteps for 一键生成全书."""
    vol_dirs = sorted(
        [d for d in novel_dir.iterdir() if d.is_dir() and d.name.startswith("volume_")],
        key=_vol_sort_key,
    )
    steps = []
    num_map = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    volumes_info = []
    for i, vol_dir in enumerate(vol_dirs):
        vol_name = meta.get("volumes", {}).get(vol_dir.name, f"第{num_map[i] if i < len(num_map) else i + 1}卷")
        chapters = sorted(vol_dir.glob("chapter_*.md"), key=_ch_sort_key)
        volumes_info.append({"dir": vol_dir.name, "name": vol_name, "chapters": [{"id": c.stem, "filename": c.name} for c in chapters]})
    steps.append({"type": "outline", "path": "global_outline.md", "label": "整体大纲", "phase": "大纲规划"})
    for vi, vol in enumerate(volumes_info):
        steps.append({"type": "vol_outline", "volIdx": vi, "path": f"{vol['dir']}/outline_detailed.md", "label": f"{vol['name']} 大纲", "phase": "大纲规划"})
    steps.append({"type": "characters", "path": "characters.md", "label": "人物设定", "phase": "人物设计"})
    for vi, vol in enumerate(volumes_info):
        for ci, ch in enumerate(vol["chapters"]):
            steps.append({"type": "chapter", "volIdx": vi, "chIdx": ci, "path": f"{vol['dir']}/{ch['filename']}", "label": ch["id"], "phase": "章节创作"})
    return steps


def _build_novel_chapters(novel_dir: Path, meta: dict):
    """Yield (volume_name, chapter_title, content) for each chapter."""
    vol_dirs = sorted(
        [d for d in novel_dir.iterdir() if d.is_dir() and d.name.startswith("volume_")],
        key=_vol_sort_key,
    )
    for vol_dir in vol_dirs:
        vol_name = meta.get("volumes", {}).get(vol_dir.name, vol_dir.name)
        for cf in sorted(vol_dir.glob("chapter_*.md"), key=_ch_sort_key):
            content = cf.read_text(encoding="utf-8")
            title = ""
            lines = content.split("\n")
            for line in lines:
                if line.strip().startswith("# "):
                    title = line.strip()[2:]
                    break
            # 正文去掉首行 # 标题（导出时单独用 title）
            body_lines = []
            skip_first_heading = True
            for line in lines:
                if skip_first_heading and line.strip().startswith("# "):
                    skip_first_heading = False
                    continue
                body_lines.append(line)
            body = "\n".join(body_lines).strip()
            yield vol_name, title or cf.stem, body


@app.route("/api/novel/export", methods=["GET"])
def novel_export():
    """导出全书为 TXT 或 Word。GET ?novel_id=xxx&format=txt|docx"""
    novel_id = request.args.get("novel_id")
    fmt = (request.args.get("format") or "txt").lower()
    if not novel_id or ".." in novel_id:
        return jsonify({"error": "Invalid novel_id"}), 400
    if fmt not in ("txt", "docx"):
        return jsonify({"error": "format must be txt or docx"}), 400
    if fmt == "docx" and not _DOCX_AVAILABLE:
        return jsonify({"error": "Word 导出需要安装 python-docx"}), 503

    novel_dir = BOOKS_ROOT / novel_id
    if not novel_dir.exists():
        return jsonify({"error": "Novel not found"}), 404

    meta = {}
    if (novel_dir / "meta.json").exists():
        meta = json.loads((novel_dir / "meta.json").read_text(encoding="utf-8"))
    title = meta.get("title", novel_id)
    safe_name = re.sub(r'[/\\?*:|"]', "_", title)[:80]

    def content_disposition(ext):
        # HTTP 头仅支持 Latin-1，中文文件名用 RFC 5987 filename*=UTF-8''
        full_name = f"{safe_name}_全文.{ext}"
        ascii_fallback = f"export.{ext}"
        encoded = quote(full_name, safe="")
        return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"

    if fmt == "txt":
        parts = [title, "=" * min(60, len(title)), ""]
        cur_vol = None
        for vol_name, ch_title, body in _build_novel_chapters(novel_dir, meta):
            if vol_name != cur_vol:
                parts.append("")
                parts.append(vol_name)
                parts.append("")
                cur_vol = vol_name
            parts.append(ch_title)
            parts.append("")
            parts.append(body)
            parts.append("")
        text = "\n".join(parts)
        return Response(
            text,
            mimetype="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": content_disposition("txt"),
            },
        )

    # docx
    doc = Document()
    doc.add_heading(title, 0)
    cur_vol = None
    for vol_name, ch_title, body in _build_novel_chapters(novel_dir, meta):
        if vol_name != cur_vol:
            doc.add_paragraph()
            p = doc.add_paragraph(vol_name)
            p.runs[0].bold = True
            cur_vol = vol_name
        doc.add_heading(ch_title, level=1)
        for para in body.split("\n\n"):
            para = para.replace("\n", " ").strip()
            if para:
                doc.add_paragraph(para)
    buf = __import__("io").BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": content_disposition("docx"),
        },
    )


@app.route("/api/novel/clear-generated", methods=["POST"])
def novel_clear_generated():
    """清空已生成的大纲、人物与章节正文，便于修改核心构思后「一键生成全书」从头生成。保留卷/章文件结构，章节文件写回占位。"""
    data = request.json
    novel_id = data.get("novel_id")
    if not novel_id or ".." in novel_id:
        return jsonify({"error": "Invalid novel_id"}), 400

    novel_dir = BOOKS_ROOT / novel_id
    if not novel_dir.exists():
        return jsonify({"error": "Novel not found"}), 404

    removed = []
    # 删除整体大纲、人物设定、挖坑填坑记录
    for name in ("global_outline.md", "characters.md", "plot_threads.md"):
        p = novel_dir / name
        if p.exists():
            p.unlink()
            removed.append(name)

    vol_dirs = sorted(
        [d for d in novel_dir.iterdir() if d.is_dir() and d.name.startswith("volume_")],
        key=_vol_sort_key,
    )
    num_map = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    for i, vol_dir in enumerate(vol_dirs):
        # 删除卷大纲
        outline_p = vol_dir / "outline_detailed.md"
        if outline_p.exists():
            outline_p.unlink()
            removed.append(f"{vol_dir.name}/outline_detailed.md")
        # 章节文件写回占位，保证结构不变且生成时不跳过
        for cf in sorted(vol_dir.glob("chapter_*.md"), key=_ch_sort_key):
            m = re.match(r"chapter_(\d+)", cf.stem)
            num = m.group(1) if m else cf.stem
            cf.write_text(f"# 第{num}章\n", encoding="utf-8")
            removed.append(f"{vol_dir.name}/{cf.name}（已清空）")

    # 将 meta.json 中的卷标题恢复为默认「第一卷」「第二卷」，清空大纲/生成时写的卷名
    meta_path = novel_dir / "meta.json"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        volumes = meta.get("volumes") or {}
        for i, vol_dir in enumerate(vol_dirs):
            volumes[vol_dir.name] = f"第{num_map[i] if i < len(num_map) else i + 1}卷"
        meta["volumes"] = volumes
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
        removed.append("meta.json（卷标题已恢复为默认）")

    return jsonify({"ok": True, "removed": removed})


@app.route("/api/novel/read", methods=["POST"])
def novel_read():
    data = request.json
    novel_id = data["novel_id"]
    file_path = data["path"]

    if ".." in file_path:
        return jsonify({"error": "Invalid path"}), 400

    filepath = BOOKS_ROOT / novel_id / file_path
    if not filepath.exists():
        return jsonify({"content": "", "exists": False})

    content = filepath.read_text(encoding="utf-8")
    return jsonify({"content": content, "exists": True})


@app.route("/api/novel/save", methods=["POST"])
def novel_save():
    data = request.json
    novel_id = data["novel_id"]
    file_path = data["path"]
    content = data["content"]

    if ".." in file_path:
        return jsonify({"error": "Invalid path"}), 400

    filepath = BOOKS_ROOT / novel_id / file_path
    filepath.parent.mkdir(parents=True, exist_ok=True)
    filepath.write_text(content, encoding="utf-8")

    return jsonify({"ok": True})


# ──────────────────────────── Non-streaming AI (for utility tasks) ────────────────────────────

@app.route("/api/ai/complete", methods=["POST"])
def ai_complete():
    data = request.json
    api_key, base_url, model = resolve_ai_config(data)
    messages = data.get("messages", [])

    if not api_key:
        return jsonify({"error": "未配置 API Key"}), 400

    url = f"{base_url}/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "max_tokens": data.get("max_tokens", 16384)}

    try:
        resp = http_requests.post(url, json=payload, headers=headers, timeout=300)
        if resp.status_code != 200:
            return jsonify({"error": f"API 错误 ({resp.status_code}): {resp.text[:500]}"}), 500
        result = resp.json()
        content = result["choices"][0]["message"].get("content") or ""
        # Strip <think>...</think> blocks from reasoning models
        content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
        if not content:
            return jsonify({"error": "模型未返回有效内容，可能需要增加 max_tokens"}), 500
        return jsonify({"content": content})
    except http_requests.exceptions.ConnectionError:
        return jsonify({"error": "无法连接到 AI 服务"}), 500
    except http_requests.exceptions.Timeout:
        return jsonify({"error": "AI 服务响应超时"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ──────────────────────────── Consistency System ────────────────────────────

@app.route("/api/novel/save-facts", methods=["POST"])
def save_facts():
    data = request.json
    novel_id = data["novel_id"]
    chapter_id = data["chapter_id"]
    facts = data["facts"]

    consistency_dir = BOOKS_ROOT / novel_id / "consistency"
    consistency_dir.mkdir(parents=True, exist_ok=True)
    filepath = consistency_dir / f"{chapter_id}_facts.json"
    filepath.write_text(json.dumps(facts, ensure_ascii=False, indent=2), encoding="utf-8")
    return jsonify({"ok": True})


@app.route("/api/novel/save-summary", methods=["POST"])
def save_summary():
    data = request.json
    novel_id = data["novel_id"]
    chapter_id = data["chapter_id"]
    summary = data["summary"]

    consistency_dir = BOOKS_ROOT / novel_id / "consistency"
    consistency_dir.mkdir(parents=True, exist_ok=True)
    summaries_path = consistency_dir / "summaries.json"

    summaries = {}
    if summaries_path.exists():
        try:
            summaries = json.loads(summaries_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass
    summaries[chapter_id] = summary
    summaries_path.write_text(json.dumps(summaries, ensure_ascii=False, indent=2), encoding="utf-8")
    return jsonify({"ok": True})


@app.route("/api/novel/load-summaries", methods=["POST"])
def load_summaries():
    data = request.json
    novel_id = data["novel_id"]

    summaries_path = BOOKS_ROOT / novel_id / "consistency" / "summaries.json"
    if not summaries_path.exists():
        return jsonify({"summaries": {}})
    try:
        summaries = json.loads(summaries_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        summaries = {}
    return jsonify({"summaries": summaries})


@app.route("/api/novel/check-consistency", methods=["POST"])
def check_consistency():
    data = request.json
    novel_id = data["novel_id"]
    novel_dir = BOOKS_ROOT / novel_id

    if not novel_dir.exists():
        return jsonify({"error": "Novel not found"}), 404

    report = run_consistency_check(novel_dir)
    return jsonify(report)


@app.route("/api/novel/update-meta", methods=["POST"])
def novel_update_meta():
    data = request.json
    novel_id = data["novel_id"]
    updates = data["updates"]

    meta_path = BOOKS_ROOT / novel_id / "meta.json"
    if not meta_path.exists():
        return jsonify({"error": "meta.json not found"}), 404

    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    meta.update(updates)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    return jsonify({"ok": True})


# ──────────────────────────── Background Tasks ────────────────────────────

@app.route("/api/task/start", methods=["POST"])
def task_start():
    data = request.json
    novel_id = data["novel_id"]
    task_type = data.get("task_type", "auto")   # 'single' or 'auto'
    items = data.get("items", [])

    # Resolve AI config (user-provided or built-in)
    ai_cfg = resolve_ai_config(data)
    ai_config = {"api_key": ai_cfg[0], "base_url": ai_cfg[1], "model": ai_cfg[2]}

    result = task_manager.start_task(
        novel_id, task_type, items, BOOKS_ROOT, ai_config
    )
    return jsonify(result)


@app.route("/api/task/status/<novel_id>", methods=["GET"])
def task_status(novel_id):
    return jsonify(task_manager.get_status(novel_id))


@app.route("/api/task/log/<novel_id>", methods=["GET"])
def task_log(novel_id):
    since = request.args.get("since", 0, type=int)
    entries = task_manager.get_log(novel_id, since)
    return jsonify({"entries": entries, "next_since": since + len(entries)})


@app.route("/api/task/stop/<novel_id>", methods=["POST"])
def task_stop(novel_id):
    return jsonify(task_manager.stop_task(novel_id))


# ──────────────────────────── Publish / Unpublish ────────────────────────────

@app.route("/api/novel/unpublish", methods=["POST"])
def novel_unpublish():
    """Mark a novel as unpublished and rebuild the site."""
    data = request.json
    novel_id = data["novel_id"]

    meta_path = BOOKS_ROOT / novel_id / "meta.json"
    if not meta_path.exists():
        return jsonify({"error": "meta.json not found"}), 404

    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    meta["published"] = False
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    # Remove published data directory
    import shutil
    data_dir = WEBSITE_DIR / "data" / novel_id
    if data_dir.exists():
        shutil.rmtree(data_dir)

    # Rebuild catalog
    result = subprocess.run(
        ["python3", "build_site.py"],
        cwd=str(WEBSITE_DIR),
        capture_output=True, text=True,
    )
    return jsonify({"ok": True, "output": result.stdout})


import random
import threading

_cover_tasks = {}  # novel_id -> {status, progress, error, prompt_used, cover_url}

_COVER_SYSTEM_PROMPTS = [
    (
        "You are a professional book cover designer. Based on the novel info provided, "
        "generate a single concise English prompt for an AI image generator to create a "
        "stunning book cover illustration. Requirements:\n"
        "- Characters should be depicted as artistic silhouettes, distant figures, or "
        "shown from behind — avoid detailed faces or close-up portraits\n"
        "- Include mood, color palette, and artistic style (e.g. digital painting, anime style, "
        "concept art, cinematic illustration)\n"
        "- Do NOT include any text, titles, letters, or words in the image\n"
        "- Keep it elegant and family-friendly, no violence, blood, or nudity\n"
        "- Combine character silhouettes with epic scenery or atmospheric backgrounds\n"
        "Output ONLY the English prompt, nothing else. Keep it under 200 words."
    ),
    (
        "You are a professional book cover artist. Create an English prompt for an AI image "
        "generator. The image should be a beautiful, atmospheric digital painting suitable "
        "for a novel cover. Requirements:\n"
        "- Show a dramatic landscape or cityscape that reflects the story's world\n"
        "- You may include one or two distant human silhouettes seen from behind or far away\n"
        "- Use rich, vivid colors and cinematic lighting\n"
        "- Absolutely NO text, words, letters, or symbols in the image\n"
        "- Absolutely NO close-up faces, NO weapons, NO blood, NO nudity\n"
        "- Style: concept art, matte painting, or fantasy illustration\n"
        "Output ONLY the English prompt, nothing else. Keep it under 150 words."
    ),
    (
        "You are a book cover designer. Generate a safe English prompt for an AI image "
        "generator. The image must be a pure landscape or abstract scene:\n"
        "- Beautiful fantasy/sci-fi landscape, sky, mountains, ocean, forest, or city\n"
        "- NO people, NO faces, NO human figures at all\n"
        "- NO text, NO letters, NO words\n"
        "- Rich colors, dramatic lighting, digital art style\n"
        "- Must be completely safe and family-friendly\n"
        "Output ONLY the English prompt. Keep it under 100 words."
    ),
]


def _run_cover_task(novel_id, ai_config, api_key):
    """Background worker for cover generation."""
    task = _cover_tasks[novel_id]
    meta_path = BOOKS_ROOT / novel_id / "meta.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    title = meta.get("title", "")
    subtitle = meta.get("subtitle", "")
    description = meta.get("description", "") or meta.get("premise", "")
    tags = ", ".join(meta.get("tags", []))

    text_api_key, text_base_url, text_model = ai_config
    user_msg = f"Novel title: {title}\nSubtitle: {subtitle}\nDescription: {description}\nTags/Genre: {tags}"

    flux_headers = {
        "Authorization": f"Bearer {api_key}",
        "Accept": "application/json",
        "Content-Type": "application/json",
    }

    for attempt, sys_prompt in enumerate(_COVER_SYSTEM_PROMPTS):
        if task.get("cancelled"):
            task["status"] = "cancelled"
            return

        # Step 1: Generate image prompt
        task["progress"] = f"正在生成封面描述（第{attempt + 1}次）…"
        try:
            url = f"{text_base_url}/chat/completions"
            headers = {"Authorization": f"Bearer {text_api_key}", "Content-Type": "application/json"}
            resp = http_requests.post(url, json={
                "model": text_model,
                "messages": [
                    {"role": "system", "content": sys_prompt},
                    {"role": "user", "content": user_msg},
                ],
                "max_tokens": 8192,
            }, headers=headers, timeout=300)
            if resp.status_code != 200:
                task["progress"] = f"AI 描述生成失败 ({resp.status_code})，重试中…"
                continue
            result = resp.json()
            image_prompt = result["choices"][0]["message"].get("content") or ""
            image_prompt = re.sub(r"<think>.*?</think>", "", image_prompt, flags=re.DOTALL).strip()
            if not image_prompt:
                task["progress"] = "AI 返回空内容，重试中…"
                continue
        except Exception as e:
            task["progress"] = f"描述生成异常: {str(e)[:100]}，重试中…"
            continue

        # Step 2: Generate image with FLUX (try 2 seeds per prompt)
        for seed_try in range(2):
            if task.get("cancelled"):
                task["status"] = "cancelled"
                return

            task["progress"] = f"正在生成封面图片（第{attempt + 1}次，种子{seed_try + 1}）…"
            try:
                flux_payload = {
                    "prompt": image_prompt,
                    "mode": "base",
                    "cfg_scale": 5,
                    "width": 768,
                    "height": 1024,
                    "seed": random.randint(1, 2**32 - 1),
                    "steps": 50,
                }
                flux_resp = http_requests.post(FLUX_IMAGE_URL, json=flux_payload,
                                                headers=flux_headers, timeout=180)
                if flux_resp.status_code != 200:
                    continue

                flux_result = flux_resp.json()
                artifacts = flux_result.get("artifacts", [])
                if not artifacts:
                    continue

                finish = artifacts[0].get("finishReason", "UNKNOWN")
                if finish == "CONTENT_FILTERED":
                    task["progress"] = "内容被过滤，正在使用更安全的描述重试…"
                    break  # break seed loop, try next prompt
                if finish != "SUCCESS":
                    continue

                # Success!
                img_bytes = base64.b64decode(artifacts[0]["base64"])
                cover_path = BOOKS_ROOT / novel_id / "cover.jpg"
                cover_path.write_bytes(img_bytes)

                meta["cover_image"] = True
                meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

                task["status"] = "done"
                task["progress"] = "封面生成成功！"
                task["prompt_used"] = image_prompt
                task["cover_url"] = f"/api/novel/cover/{novel_id}?t={int(datetime.now().timestamp())}"
                return
            except Exception:
                continue

    task["status"] = "error"
    task["error"] = "封面生成失败（已用尽所有策略），请重试"


@app.route("/api/novel/generate-cover", methods=["POST"])
def generate_cover():
    """Start async cover generation."""
    data = request.json
    novel_id = data["novel_id"]
    meta_path = BOOKS_ROOT / novel_id / "meta.json"
    if not meta_path.exists():
        return jsonify({"error": "meta.json not found"}), 404

    existing = _cover_tasks.get(novel_id)
    if existing and existing.get("status") == "running":
        return jsonify({"ok": True, "status": "running", "progress": existing.get("progress", "")})

    ai_config = resolve_ai_config(data)
    api_key = data.get("api_key") or os.getenv("AI_API_KEY") or BUILTIN_API_KEY

    task = {"status": "running", "progress": "正在启动…", "error": None, "prompt_used": None, "cover_url": None}
    _cover_tasks[novel_id] = task

    t = threading.Thread(target=_run_cover_task, args=(novel_id, ai_config, api_key), daemon=True)
    t.start()

    return jsonify({"ok": True, "status": "running"})


@app.route("/api/novel/cover-status/<novel_id>")
def cover_status(novel_id):
    """Poll cover generation status."""
    task = _cover_tasks.get(novel_id)
    if not task:
        return jsonify({"status": "idle"})
    return jsonify(task)


@app.route("/api/novel/cover/<novel_id>")
def serve_cover(novel_id):
    """Serve the generated cover image."""
    cover_path = BOOKS_ROOT / novel_id / "cover.jpg"
    if not cover_path.exists():
        return "", 404
    return send_from_directory(str(cover_path.parent), cover_path.name, mimetype="image/jpeg")


_INTRO_SYSTEM = """你是一位网文运营专家。请根据小说的大纲与标题，写一段面向读者的「小说介绍」。
要求：
1. 2～4 句话，简练有悬念，吸引读者点开阅读。
2. 不要复述核心构思或剧透结局，不要写成「本书讲述了…」的说明书式。
3. 用情境、冲突或悬念开头，让读者好奇接下来会发生什么。
只输出介绍正文，不要标题、不要引号、不要「介绍：」等前缀。"""


@app.route("/api/novel/generate-intro", methods=["POST"])
def generate_intro():
    """AI 根据大纲生成简练有悬念的小说介绍，写入 meta.description。"""
    data = request.json
    novel_id = data.get("novel_id")
    if not novel_id or ".." in novel_id:
        return jsonify({"error": "Invalid novel_id"}), 400

    novel_dir = BOOKS_ROOT / novel_id
    meta_path = novel_dir / "meta.json"
    if not novel_dir.exists() or not meta_path.exists():
        return jsonify({"error": "Novel not found"}), 404

    outline_path = novel_dir / "global_outline.md"
    outline = outline_path.read_text(encoding="utf-8")[:2500] if outline_path.exists() else ""
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    title = meta.get("title", "")
    premise = meta.get("premise") or meta.get("description", "")

    api_key, base_url, model = resolve_ai_config(data)
    if not api_key:
        return jsonify({"error": "未配置 API Key"}), 400

    user_msg = f"书名：《{title}》\n\n核心构思（仅供参考，不要照抄）：\n{premise[:800]}\n\n全局大纲（节选）：\n{outline}"
    try:
        url = f"{base_url.rstrip('/')}/chat/completions"
        resp = http_requests.post(
            url,
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": _INTRO_SYSTEM},
                    {"role": "user", "content": user_msg},
                ],
                "max_tokens": 512,
            },
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            timeout=60,
        )
        if resp.status_code != 200:
            return jsonify({"error": f"AI 请求失败 ({resp.status_code})"}), 502
        result = resp.json()
        intro = (result.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        intro = re.sub(r"<think>.*?</think>", "", intro, flags=re.DOTALL).strip()
        intro = intro.strip('"\'「」').strip()
        if not intro:
            return jsonify({"error": "AI 未返回有效介绍"}), 502
        meta["description"] = intro
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
        return jsonify({"ok": True, "description": intro})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ──────────────────────────── Auto-seed & Pipeline ────────────────────────────

_AUTO_SEED_SYSTEM = """你是一位网文策划。请根据当前网文流行趋势与常见题材，为指定频道生成一部新小说的完整配置。
只输出一个 JSON 对象，不要其他文字。字段如下（全部必填）：
- premise: 核心构思（2～5 句话，含主角设定、主线、风格）
- title: 中文书名（2～5 字）
- subtitle: 英文副标题
- tags: 标签数组，如 ["玄幻", "热血", "升级流"]
- writing_style: 写作风格，如 "天蚕土豆" 或 "轻松幽默、节奏快"
- volume_count: 卷数（3～5 的整数）
- total_chapters: 总章数（50～150 的整数）"""


def _do_auto_seed(channel: str, api_key: str, base_url: str, model: str) -> dict:
    """Internal: AI 生成种子配置，返回 seed 字典。"""
    channel = "male" if channel not in ("male", "female") else channel
    channel_cn = "男频" if channel == "male" else "女频"
    url = f"{base_url.rstrip('/')}/chat/completions"
    resp = http_requests.post(
        url,
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": _AUTO_SEED_SYSTEM},
                {"role": "user", "content": f"频道：{channel_cn}。请生成一部适合该频道的原创小说配置。"},
            ],
            "max_tokens": 1024,
        },
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        timeout=90,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"AI 请求失败 ({resp.status_code})")
    result = resp.json()
    raw = (result.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
    raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
    m = re.search(r"\{[\s\S]*\}", raw)
    if not m:
        raise RuntimeError("AI 未返回有效 JSON")
    seed = json.loads(m.group(0))
    seed.setdefault("premise", "")
    seed.setdefault("title", "未命名")
    seed.setdefault("subtitle", "")
    seed.setdefault("tags", [])
    seed.setdefault("writing_style", "")
    seed.setdefault("volume_count", 3)
    seed.setdefault("total_chapters", 50)
    seed["channel"] = channel
    if isinstance(seed.get("tags"), str):
        seed["tags"] = [t.strip() for t in seed["tags"].split(",") if t.strip()]
    seed["volume_count"] = max(1, min(10, int(seed["volume_count"])))
    seed["total_chapters"] = max(10, min(500, int(seed["total_chapters"])))
    return seed


@app.route("/api/novel/auto-seed", methods=["POST"])
def auto_seed():
    """AI 根据男频/女频生成一部新小说的核心构思与基础配置（无需人工提供构思）。"""
    data = request.json or {}
    channel = (data.get("channel") or "male").lower()
    api_key, base_url, model = resolve_ai_config(data)
    if not api_key:
        return jsonify({"error": "未配置 API Key"}), 400
    try:
        seed = _do_auto_seed(channel, api_key, base_url, model)
        return jsonify({"ok": True, "seed": seed})
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON 解析失败: {e}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _do_init_novel(meta: dict, volume_count: int, total_chapters: int) -> str:
    """Internal: 创建小说目录与文件，返回 novel_id。"""
    title = meta.get("title", "novel")
    novel_id = _title_to_dir_id(title)
    base_id = novel_id
    counter = 2
    while (BOOKS_ROOT / novel_id).exists():
        novel_id = f"{base_id}_{counter}"
        counter += 1
    novel_dir = BOOKS_ROOT / novel_id
    novel_dir.mkdir(parents=True)
    base = total_chapters // volume_count
    remainder = total_chapters % volume_count
    chapters_per_vol = [base + (1 if i < remainder else 0) for i in range(volume_count)]
    volumes = {}
    num_map = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    ch_counter = 1
    for i in range(volume_count):
        vol_dir_name = f"volume_{i + 1}"
        vol_display = f"第{num_map[i] if i < len(num_map) else i + 1}卷"
        volumes[vol_dir_name] = vol_display
        vol_path = novel_dir / vol_dir_name
        vol_path.mkdir()
        for _ in range(chapters_per_vol[i]):
            (vol_path / f"chapter_{ch_counter:03d}.md").write_text(f"# 第{ch_counter:03d}章\n", encoding="utf-8")
            ch_counter += 1
    meta_data = {
        "title": meta["title"],
        "subtitle": meta.get("subtitle", ""),
        "author": meta.get("author", "AI & Thomas"),
        "description": meta.get("description", ""),
        "tags": meta.get("tags", []),
        "volumes": volumes,
    }
    if meta.get("premise") is not None:
        meta_data["premise"] = meta["premise"]
    if meta.get("writing_style"):
        meta_data["writing_style"] = meta["writing_style"]
    if meta.get("channel"):
        meta_data["channel"] = meta["channel"]
    (novel_dir / "meta.json").write_text(json.dumps(meta_data, ensure_ascii=False, indent=2), encoding="utf-8")
    (novel_dir / "story_bible.md").write_text(
        f"# {meta['title']} · 故事圣经\n\n"
        "> 本文件由系统自动维护，记录全书事实基准。AI 创作时必须严格遵守。\n\n"
        "## 角色登记\n\n| 角色 | 首次出现 | 身份 | 关键特征 |\n|------|---------|------|----------|\n\n"
        "## 时间线\n\n（暂无）\n\n## 关键事件\n\n（暂无）\n\n## 角色状态变更\n\n（暂无）\n",
        encoding="utf-8",
    )
    (novel_dir / "consistency").mkdir(exist_ok=True)
    return novel_id


def _do_generate_intro(novel_id: str, api_key: str, base_url: str, model: str) -> None:
    """Internal: 根据大纲生成小说介绍并写入 meta.description。"""
    novel_dir = BOOKS_ROOT / novel_id
    meta_path = novel_dir / "meta.json"
    outline_path = novel_dir / "global_outline.md"
    outline = outline_path.read_text(encoding="utf-8")[:2500] if outline_path.exists() else ""
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    premise = meta.get("premise") or meta.get("description", "")
    user_msg = f"书名：《{meta.get('title', '')}》\n\n核心构思（仅供参考，不要照抄）：\n{premise[:800]}\n\n全局大纲（节选）：\n{outline}"
    url = f"{base_url.rstrip('/')}/chat/completions"
    resp = http_requests.post(
        url,
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": _INTRO_SYSTEM},
                {"role": "user", "content": user_msg},
            ],
            "max_tokens": 512,
        },
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        timeout=60,
    )
    if resp.status_code != 200 or not resp.json().get("choices"):
        return
    intro = (resp.json()["choices"][0].get("message", {}).get("content") or "").strip()
    intro = re.sub(r"<think>.*?</think>", "", intro, flags=re.DOTALL).strip().strip('"\'「」')
    if intro:
        meta["description"] = intro
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


_pipeline_lock = threading.Lock()
_pipeline_state = {"running": False, "stop": False, "channel": "male", "novel_id": None, "error": None}


def _run_auto_pipeline(ai_cfg: dict, api_key: str, alternate_channels: bool, auto_publish: bool):
    """Background: 连续自动生成小说（男频/女频轮流），每部：seed → init → 全书生成 → 介绍 → 封面 → 可选发布。"""
    global _pipeline_state
    base_url = ai_cfg.get("base_url") or ""
    model = ai_cfg.get("model") or ""
    channel = _pipeline_state.get("channel", "male")
    while True:
        with _pipeline_lock:
            if _pipeline_state.get("stop"):
                _pipeline_state["running"] = False
                _pipeline_state["stop"] = False
                return
            _pipeline_state["error"] = None

        # 1) Auto-seed（内部函数，不依赖 request）
        try:
            seed = _do_auto_seed(channel, api_key, base_url, model)
        except Exception as e:
            with _pipeline_lock:
                _pipeline_state["running"] = False
                _pipeline_state["error"] = str(e)
            return

        # 2) Init novel（内部函数）
        meta = {
            "title": seed["title"],
            "subtitle": seed.get("subtitle", ""),
            "author": "AI & Thomas",
            "premise": seed["premise"],
            "description": "",
            "tags": seed.get("tags", []),
            "writing_style": seed.get("writing_style", ""),
            "channel": seed.get("channel", channel),
        }
        try:
            novel_id = _do_init_novel(meta, seed["volume_count"], seed["total_chapters"])
        except Exception as e:
            with _pipeline_lock:
                _pipeline_state["running"] = False
                _pipeline_state["error"] = str(e)
            return
        with _pipeline_lock:
            _pipeline_state["novel_id"] = novel_id

        novel_dir = BOOKS_ROOT / novel_id
        meta_path = novel_dir / "meta.json"
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        steps = _build_auto_steps(novel_dir, meta)
        task_manager.start_task(novel_id, "auto", steps, BOOKS_ROOT, ai_cfg)

        # 3) Wait for task done (poll)
        while True:
            time.sleep(10)
            with _pipeline_lock:
                if _pipeline_state.get("stop"):
                    task_manager.stop_task(novel_id)
                    _pipeline_state["running"] = False
                    return
            st = task_manager.get_status(novel_id)
            if st.get("status") not in ("running",):
                break
        if st.get("status") != "done":
            with _pipeline_lock:
                _pipeline_state["running"] = False
                _pipeline_state["error"] = st.get("error") or "任务未完成"
            return

        # 4) Generate intro（内部函数）
        try:
            _do_generate_intro(novel_id, api_key, base_url, model)
        except Exception:
            pass

        # 5) Cover (sync run in thread)
        _run_cover_task(novel_id, (ai_cfg.get("api_key"), ai_cfg.get("base_url"), ai_cfg.get("model")), api_key)

        # 6) Optional publish
        if auto_publish:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            meta["published"] = True
            meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
            subprocess.run(["python3", "build_site.py"], cwd=str(WEBSITE_DIR), capture_output=True, timeout=60)

        # Next channel
        if alternate_channels:
            channel = "female" if channel == "male" else "male"
        with _pipeline_lock:
            _pipeline_state["channel"] = channel
            if _pipeline_state.get("stop"):
                _pipeline_state["running"] = False
                return


@app.route("/api/novel/auto-pipeline/start", methods=["POST"])
def auto_pipeline_start():
    """启动连续自动生产：男频/女频轮流生成小说，每部自动完成构思→全书→介绍→封面。"""
    data = request.json or {}
    with _pipeline_lock:
        if _pipeline_state["running"]:
            return jsonify({"ok": False, "error": "流水线已在运行"})
        _pipeline_state["running"] = True
        _pipeline_state["stop"] = False
        _pipeline_state["error"] = None
        _pipeline_state["channel"] = data.get("channel") or "male"
        _pipeline_state["novel_id"] = None

    ai_cfg = resolve_ai_config(data)
    api_key = data.get("api_key") or os.getenv("AI_API_KEY") or BUILTIN_API_KEY
    ai_config = {"api_key": ai_cfg[0], "base_url": ai_cfg[1], "model": ai_cfg[2]}
    alternate = data.get("alternate_channels", True)
    auto_publish = data.get("auto_publish", False)

    def run():
        _run_auto_pipeline(ai_config, api_key, alternate, auto_publish)

    t = threading.Thread(target=run, daemon=True)
    t.start()
    return jsonify({"ok": True, "message": "自动生产已启动（男频/女频轮流）"})


@app.route("/api/novel/auto-pipeline/stop", methods=["POST"])
def auto_pipeline_stop():
    with _pipeline_lock:
        _pipeline_state["stop"] = True
    return jsonify({"ok": True})


@app.route("/api/novel/auto-pipeline/status", methods=["GET"])
def auto_pipeline_status():
    with _pipeline_lock:
        s = dict(_pipeline_state)
    return jsonify(s)


@app.route("/api/rebuild", methods=["POST"])
def rebuild():
    result = subprocess.run(
        ["python3", "build_site.py"],
        cwd=str(WEBSITE_DIR),
        capture_output=True,
        text=True,
    )
    return jsonify({
        "ok": result.returncode == 0,
        "output": result.stdout,
        "error": result.stderr,
    })


COMMENTS_DIR = WEBSITE_DIR / "comments"


def _load_comments(filepath: Path) -> list:
    if filepath.exists():
        try:
            return json.loads(filepath.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            pass
    return []


def _save_comments(filepath: Path, comments: list):
    filepath.parent.mkdir(parents=True, exist_ok=True)
    filepath.write_text(json.dumps(comments, ensure_ascii=False, indent=2), encoding="utf-8")


# ──────────────────────────── Chapter Comments ────────────────────────────

@app.route("/api/comments/chapter/<book_id>/<chapter_id>", methods=["GET"])
def get_chapter_comments(book_id, chapter_id):
    filepath = COMMENTS_DIR / book_id / f"{chapter_id}.json"
    return jsonify(_load_comments(filepath))


@app.route("/api/comments/chapter", methods=["POST"])
def add_chapter_comment():
    data = request.json
    book_id = data["book_id"]
    chapter_id = data["chapter_id"]
    comment = {
        "id": str(uuid.uuid4())[:8],
        "author": data.get("author", "").strip() or "匿名读者",
        "text": data["text"].strip(),
        "quote": data.get("quote", "").strip(),
        "paragraph_index": data.get("paragraph_index"),
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }
    if not comment["text"]:
        return jsonify({"error": "评论内容不能为空"}), 400

    filepath = COMMENTS_DIR / book_id / f"{chapter_id}.json"
    comments = _load_comments(filepath)
    comments.append(comment)
    _save_comments(filepath, comments)
    return jsonify({"ok": True, "comment": comment})


@app.route("/api/comments/chapter/<book_id>/<chapter_id>/<comment_id>", methods=["DELETE"])
def delete_chapter_comment(book_id, chapter_id, comment_id):
    filepath = COMMENTS_DIR / book_id / f"{chapter_id}.json"
    comments = _load_comments(filepath)
    comments = [c for c in comments if c["id"] != comment_id]
    _save_comments(filepath, comments)
    return jsonify({"ok": True})


# ──────────────────────────── Book Reviews ────────────────────────────

@app.route("/api/comments/book/<book_id>", methods=["GET"])
def get_book_reviews(book_id):
    filepath = COMMENTS_DIR / f"book_{book_id}.json"
    return jsonify(_load_comments(filepath))


@app.route("/api/comments/book", methods=["POST"])
def add_book_review():
    data = request.json
    book_id = data["book_id"]
    review = {
        "id": str(uuid.uuid4())[:8],
        "author": data.get("author", "").strip() or "匿名读者",
        "text": data["text"].strip(),
        "rating": min(10, max(1, int(data.get("rating", 10)))),
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }
    if not review["text"]:
        return jsonify({"error": "评论内容不能为空"}), 400

    filepath = COMMENTS_DIR / f"book_{book_id}.json"
    reviews = _load_comments(filepath)
    reviews.append(review)
    _save_comments(filepath, reviews)
    return jsonify({"ok": True, "review": review})


@app.route("/api/comments/book/<book_id>/<review_id>", methods=["DELETE"])
def delete_book_review(book_id, review_id):
    filepath = COMMENTS_DIR / f"book_{book_id}.json"
    reviews = _load_comments(filepath)
    reviews = [r for r in reviews if r["id"] != review_id]
    _save_comments(filepath, reviews)
    return jsonify({"ok": True})


# ──────────────────────────── Product Comments ────────────────────────────

@app.route("/api/comments/product", methods=["GET"])
def get_product_comments():
    filepath = COMMENTS_DIR / "product.json"
    return jsonify(_load_comments(filepath))


@app.route("/api/comments/product", methods=["POST"])
def add_product_comment():
    data = request.json
    comment = {
        "id": str(uuid.uuid4())[:8],
        "author": data.get("author", "").strip() or "匿名用户",
        "text": data["text"].strip(),
        "rating": min(5, max(1, int(data.get("rating", 5)))),
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }
    if not comment["text"]:
        return jsonify({"error": "评论内容不能为空"}), 400

    filepath = COMMENTS_DIR / "product.json"
    comments = _load_comments(filepath)
    comments.append(comment)
    _save_comments(filepath, comments)
    return jsonify({"ok": True, "comment": comment})


@app.route("/api/comments/product/<comment_id>", methods=["DELETE"])
def delete_product_comment(comment_id):
    filepath = COMMENTS_DIR / "product.json"
    comments = _load_comments(filepath)
    comments = [c for c in comments if c["id"] != comment_id]
    _save_comments(filepath, comments)
    return jsonify({"ok": True})


if __name__ == "__main__":
    print(f"Books root: {BOOKS_ROOT}")
    print(f"Website dir: {WEBSITE_DIR}")
    print("Starting server on http://localhost:8080")
    app.run(host="0.0.0.0", port=8080, debug=False)
